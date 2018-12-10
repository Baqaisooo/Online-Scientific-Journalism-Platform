from django.shortcuts import render, redirect
from django.http import JsonResponse
from passlib.hash import pbkdf2_sha256
from .data_info import *
from .alert_info import *
from .models import User, Writer, Article, Multimedia, Comment, Rate, Setting, Package, Subscriber, Subscription, Feedback, Purchase, Associate_Editor, Editor_in_Chief, Article_Revision, Reviewer
from django.core.mail import EmailMessage
from django.contrib import messages
from django.db.models import Avg # to get average
from django.db.models import Sum # to get summition
from django.db.models import Q  # to be able to use AND,OR operator
from django.db.models import Count
import random
import logging

import os
import shutil
import datetime
import convertapi
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from zipfile import ZipFile
from django.http import HttpResponse
from django.conf import settings


def change_lang(request, lang):
    request.session['lang'] = lang
    return redirect(mainpage)

###
# this def will load mainpage of the project
def mainpage(request):
    if request.session.has_key('role'):
        return{
            'writer' : redirect('writer_homepage'),
            'subscriber' : redirect('subscriber_homepage'),
            'EIC' : redirect('EIC_homepage'),
            'AE' : redirect('AE_homepage'),
            'reviewer' : redirect('reviewer_homepage'),
        }[ request.session['role'] ]

    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'

    #calling this function to check and update the expiered ONLINE or EDITORIAL articles to ARCHIVE
    check_and_update()
    
    latest_articles = get_16_articles('latest')            # latest means latest articles
    highest_views_artilces = get_16_articles('h_views')    # h_views means highest views articles
    highest_rates_articles = get_16_articles('h_rates')    # h_rates means highest rates articles

    context = {**reader_base_info[ request.session['lang'] ],**home_articles_info[ request.session['lang'] ]}

    context['extends'] = 'scientificJournalism/base_component/reader.html'

    context['editorial_articles'] = Article.objects.filter(Q(is_editorial=True) & Q(article_status_en='EDITORIAL')).order_by('-publishing_date')

    context['first_latest_articles'] = latest_articles.pop(0) if len(latest_articles) else [] 
    context['latest_articles'] = latest_articles

    context['first_highest_views_artilces'] = highest_views_artilces.pop(0) if len(highest_views_artilces) else [] 
    context['highest_views_artilces'] = highest_views_artilces


    context['first_highest_rates_articles'] = highest_rates_articles.pop(0) if len(highest_rates_articles) else [] 
    context['highest_rates_articles'] = highest_rates_articles
    
    return render(request,'scientificJournalism/body_component/home_articles.html',context)


# except highest articles rates 12 articles divided to 4 packages
def get_16_articles(catigoray):
    articles = ''
    if catigoray == 'latest':
        articles = Article.objects.filter(article_status_en__in=['ONLINE','EDITORIAL']).order_by('-publishing_date')[:16]
    
    elif catigoray == 'h_views':
        articles = Article.objects.filter(article_status_en__in=['ONLINE','EDITORIAL']).order_by('-number_of_views')[:16]
    
    elif catigoray == 'h_rates':
        articles = Rate.objects.values('article_ID').annotate(rating_avg = Avg('subscriber_rating')).filter(article_ID__article_status_en__in=['ONLINE','EDITORIAL']).order_by('-rating_avg')

    if(catigoray == 'latest' or catigoray == 'h_views'):
        counter = 0
        divided_articles = []
        sub_articles = []
        for article in articles:
            
            counter += 1
            comment = Comment.objects.filter(article_ID=article)
            rate = Rate.objects.filter(article_ID=article).aggregate(Avg('subscriber_rating'))
            article_package = {'article': article,'comments': len(comment),'rates': rate}
            sub_articles.append(article_package)
            
            if  ((not counter % 4) and counter):
                divided_articles.append({'packages':sub_articles,'index':counter//4})
                sub_articles = []

        if counter%4 and counter:
            divided_articles.append({'packages':sub_articles,'index':counter//4+1})
    else:
        counter = 0
        divided_articles = []
        sub_articles = []
        for article in articles:
            
            counter += 1
            comment = Comment.objects.filter(article_ID=article['article_ID'])
            article_package = {'article': Article.objects.get(article_ID = article['article_ID']),'comments': len(comment),'rates': article['rating_avg']}
            sub_articles.append(article_package)

            if  ((not counter % 3) and counter):
                divided_articles.append({'packages':sub_articles,'index':counter//3})
                sub_articles = []

        if counter%3 and counter:
            divided_articles.append({'packages':sub_articles,'index':counter//3+1})

    return divided_articles


# this def will all tring to sign in user depending on selected role
# admin can't sign in from same page that other users do .. he has a spiciel page to sign in but will keep using this def.
#######
def logging_in(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        try:
            user = User.objects.get(email=email)
        except (User.DoesNotExist):
            pass
        else:
            if user.activation != 'active':
                messages.warning(request, not_activation_alert[request.session['lang']])
                return send_activation_code(request, user.email)
            role = request.POST.get('account_role')
            
            if role == 'writer':
                try:
                    writer = Writer.objects.get(writer_ID=user.user_ID)
                except (Writer.DoesNotExist):
                    pass
                else:
                    if writer.account_status_en == 'freeze' or writer.account_status_ar == 'مجمد':
                        messages.warning(request, account_frozen_alert[request.session['lang']])
                        return redirect('mainpage') 
                    if writer.verify_password(request.POST.get('password')):
                        request.session['login'] = True
                        request.session['email'] = email
                        request.session['role'] = role
                        request.session['field'] = writer.specialty_ar
                        request.session['name'] = user.first_name+' '+ user.last_name
                        request.session['image'] = writer.picture.url
                        request.session['id'] = user.user_ID
                        return redirect('writer_homepage')

            elif role == 'subscriber':
                try:
                    subscriber = Subscriber.objects.get(subscriber_ID=user.user_ID)
                except (Subscriber.DoesNotExist):
                    pass
                else:
                    if subscriber.verify_password(request.POST.get('password')):
                        request.session['login'] = True
                        request.session['email'] = email
                        request.session['role'] = role
                        request.session['name'] = user.first_name+' '+ user.last_name
                        request.session['image'] = subscriber.picture.url
                        request.session['id'] = user.user_ID
                        
                        return redirect('subscriber_homepage')
            
            elif role == 'EIC':
                try:
                    EIC = Editor_in_Chief.objects.get(EIC_ID__EBM_ID=user)
                except (Editor_in_Chief.DoesNotExist):
                    pass
                else:
                    if EIC.EIC_ID.verify_password(request.POST.get('password')):
                        request.session['login'] = True
                        request.session['email'] = email
                        request.session['role'] = role
                        request.session['name'] = user.first_name+' '+ user.last_name
                        request.session['id'] = user.user_ID
                        
                        return redirect('EIC_homepage')
            
            elif role == 'AE':
                try:
                    AE = Associate_Editor.objects.get(AE_ID__EBM_ID=user)
                except (Associate_Editor.DoesNotExist):
                    pass
                else:
                    if AE.AE_ID.verify_password(request.POST.get('password')):
                        request.session['login'] = True
                        request.session['email'] = email
                        request.session['role'] = role
                        request.session['name'] = user.first_name+' '+ user.last_name
                        request.session['id'] = user.user_ID
                        
                        return redirect('AE_homepage')
            
            elif role == 'reviewer':
                try:
                    reviewer = Reviewer.objects.get(reviewer_ID=user)
                except (Reviewer.DoesNotExist):
                    pass
                else:
                    if reviewer.verify_password(request.POST.get('password')):
                        request.session['login'] = True
                        request.session['email'] = email
                        request.session['role'] = role
                        request.session['name'] = user.first_name+' '+ user.last_name
                        request.session['id'] = user.user_ID
                        
                        return redirect('reviewer_homepage')
 

        messages.error(request,login_error_alert[request.session['lang']])
        return redirect('mainpage') 

    return redirect('mainpage')              
                
        
#######
def writer_homepage(request):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'
        
    if(not request.session.has_key('role') or request.session['role'] != 'writer'):
        return redirect('mainpage')

    #calling this function to check and update the expiered ONLINE or EDITORIAL articles to ARCHIVE
    check_and_update()
    
    context = {
        'extends':'scientificJournalism/base_component/writer.html'
        }

    latest_articles = get_16_articles('latest')            # latest means latest articles
    highest_views_artilces = get_16_articles('h_views')    # h_views means highest views articles
    highest_rates_articles = get_16_articles('h_rates')    # h_rates means highest rates articles


    context['editorial_articles'] = Article.objects.filter(Q(is_editorial=True) & Q(article_status_en='EDITORIAL'))

    context['first_latest_articles'] = latest_articles.pop(0) if len(latest_articles) else [] 
    context['latest_articles'] = latest_articles

    context['first_highest_views_artilces'] = highest_views_artilces.pop(0) if len(highest_views_artilces) else [] 
    context['highest_views_artilces'] = highest_views_artilces


    context['first_highest_rates_articles'] = highest_rates_articles.pop(0) if len(highest_rates_articles) else [] 
    context['highest_rates_articles'] = highest_rates_articles

    context = {**context, **writer_base_info[ request.session['lang'] ],**home_articles_info[ request.session['lang'] ]}

    return render(request,'scientificJournalism/body_component/home_articles.html',context)


#######
def logout(request):
    lang = request.session['lang']
    request.session.flush()
    request.session['lang'] = lang

    messages.success(request, logout_alert[ request.session['lang'] ] )

    return redirect('mainpage')    


#######
def register(request):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'

    context = register_info[ request.session['lang'] ]
    return render(request,'scientificJournalism/one_page/register.html',context)
      

#######
def registering(request):
    if request.method == 'POST':
        try:
            user = User.objects.get(email=request.POST.get('email'))
        except User.DoesNotExist:
            pass
        else:
            messages.error(request, Email_register_alert[request.session['lang']])
            return redirect('register')
        
        user=User()
        user.first_name = request.POST.get('first_name')
        user.last_name= request.POST.get('last_name')
        user.birthdate= '-'.join(request.POST.get('birthday').split('/')[::-1])
        user.gender_en= user.gender_ar= 'MALE' if request.POST.get('gender') == 'male' else 'FEMALE'
        user.email= request.POST.get('email')
        user.phone_number= request.POST.get('phone')
        user.country= request.POST.get('country')
        user.city= request.POST.get('city')

        user.save()

        writer = Writer()
        writer.writer_ID = user
        writer.password = pbkdf2_sha256.encrypt(request.POST.get('password'), rounds=12000, salt_size=23) 
        writer.specialty_en = writer.specialty_ar = request.POST.get('specialty')
        writer.workplace = request.POST.get('workplace')
        if request.FILES.get('personal_image'):
            writer.picture = request.FILES['personal_image']
        else:
            writer.picture = 'static_image/male.png' if request.POST.get('gender') == 'male' else 'static_image/female.png'
        writer.save()
        
        return send_activation_code(request, user.email)

    else:
        return redirect('mainpage')


#######
def send_activation_code(request, to):
    """
        this function will send activation code and redirect to activation page
        INPUT:-
            to : email account for who recieve activation code.
        OUTPUT:-
            - will get activation code.
            - redirect to activation page.
    """
    activation_code = str(random.randrange(1111111,10000000))
    user = User.objects.get(email=to)
    user.activation = activation_code
    
    user.save()
    email_content = activation_email_info[ request.session['lang'] ]
    email = EmailMessage(email_content['title'], email_content['message'].format(activation_code), to=[to])
    email.send()

    email = {'email':str(to)}
    context = activation_info[ request.session['lang'] ]
    context.update(email)
    return render(request, 'scientificJournalism/one_page/account_activation.html',context)


#######
def activation(request):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'
     
    if request.method == 'POST':
        try:
            user = User.objects.get(email=request.POST.get('email'))
        except (User.DoesNotExist):
            pass
        else:
            if(user.activation is 'active'):
                return redirect('mainpage')
            if(user.activation == request.POST.get('activation_code')):
                user.activation = 'active'
                user.save()

                try:
                    writer = Writer.objects.get(writer_ID = user.user_ID)
                except (Writer.DoesNotExist):
                    pass
                else:
                    writer.account_status_en = writer.account_status_ar = 'ONLINE'
                    writer.save()
                finally:
                    messages.success(request,activation_success_alert[ request.session['lang'] ])
                # also we will do it for EIC and Reviewer .. because both have account status
            else:
                messages.error(request,activation_fail_alert[request.session['lang']])
                return send_activation_code(request, user.email)
                    
    return redirect('mainpage')


#######
def new_article(request):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'
        
    if(not request.session.has_key('role') or request.session['role'] != 'writer'):
        return redirect('mainpage')

    context = new_article_info[ request.session['lang'] ]
    return render(request,'scientificJournalism/one_page/new_article.html',context)


#######
# this function get new article data from writer to add it
def submit_article(request):
    # check if user already choose language, if not, make Arabic as a difault
    # start language checking
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'
    # end language checking
    
    # check if logged in
    if not request.session.has_key('login') or not request.session['login']:
        return redirect('mainpage')

    # check role
    if(not request.session.has_key('role') or request.session['role'] != 'writer'):
        return redirect('mainpage')

    # get submitted data
    article_field = request.POST.get('article_field')
    article_title = request.POST.get('article_title')
    article_lang = request.POST.get('article_lang')
    brief_of_article = request.POST.get('brief_of_article')
    article_context = request.POST.get('context')
    
    cover_image = request.FILES.get('cover_image')
    media_VIA = request.FILES.getlist('media')

    # get writer id from session
    writer_id = request.session['id']

    # create object from Article model and assign data to it
    # start article object
    article = Article()
    article.writer_ID = Writer.objects.get(writer_ID=writer_id)
    article.field_ar = article.field_en = article_field
    article.cover_image = cover_image
    
    if article_lang == 'ar':
        article.title_ar = article_title
        article.brief_of_article_ar = brief_of_article
        article.context_ar = article_context
    elif article_lang =='en':
        article.title_en = article_title
        article.brief_of_article_en = brief_of_article
        article.context_en = article_context

    article.save()
    # save article object

    # specify accepted extentions
    images_extentions = ['jpeg','jpg','png','gif']
    audios_extentions = ['mp3']
    videos_extentions = ['mp4']

    # set default article type 
    # article type prioritrize as [viedo, audio, image, text]
    article_type = 'TEXT'

    # go over all submitted article and save it
    for media in media_VIA:
        # create Multimedia object
        multimedia = Multimedia()
        multimedia.aritcle_ID = article
        multimedia.format = ''

        extention = str(media).split('.')[-1].lower()   # get the extention of media
        
        # check the type of extention and update the article type 
        if extention in images_extentions:
            multimedia.format = 'IMAGE'
            if article_type == 'TEXT':
                article_type = 'IMAGE'
        elif extention in videos_extentions:
            multimedia.format = 'VIDEO'
            if article_type == 'TEXT':
                article_type = 'VIDEO'
        elif extention in audios_extentions:
            multimedia.format = 'AUDIO'
            if article_type == 'TEXT':
                article_type = 'AUDIO'
        
        multimedia.content = media
        multimedia.save()
        # save multimedia object

    # add article type
    article.type = article_type
    # save article object
    article.save()

    # show successful alert message 
    messages.success(request,'تم رفع مقالتك بنجاح وهي الان تحت المراجعة، تستطيع تتبع حالة مقالتك من خلال الصفحة الحالية')
    # return to my article page to follow article status
    return redirect('my_articles')


    
def my_articles(request):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'
        
    if(not request.session.has_key('role') or request.session['role'] != 'writer'):
        return redirect('mainpage')

    my_articles = Article.objects.filter(writer_ID__writer_ID=request.session['id']).order_by('-submit_date')

    index = 1
    
    for article in my_articles:
        feedback = Feedback.objects.filter(article_ID = article.article_ID)
        if feedback.exists():
            article.feedback = feedback[0].EIC_comment
        article.index = index
        index += 1


    context = {
        'my_articles':my_articles,
        } 
    
    context = {**context, **writer_base_info[ request.session['lang'] ], **my_articles_info[ request.session['lang'] ] }

    return render(request,'scientificJournalism/body_component/my_articles.html',context)

     
#######
def subscribe(request):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'
    
    if request.session.has_key('role') and request.session['role'] != 'writer':
        return redirect('mainpage')

    context = subscribe_info[ request.session['lang'] ]
    context['packages'] = Package.objects.exclude(package_status_ar = 'INACTIVE')

    return render(request,'scientificJournalism/one_page/subscribe.html',context)


#######
def subscribing(request):
    if request.method == 'POST':
        try:
            user = User.objects.get(email=request.POST.get('email'))
        except User.DoesNotExist:
            pass
        else:
            messages.error(request, Email_register_alert[request.session['lang']])
            return redirect('subscribe')
        
        user=User()
        user.first_name = request.POST.get('first_name')
        user.last_name= request.POST.get('last_name')
        user.birthdate= '-'.join(request.POST.get('birthday').split('/')[::-1])
        user.gender_en= user.gender_ar= 'MALE' if request.POST.get('gender') == 'male' else 'FEMALE'
        user.email= request.POST.get('email')
        user.phone_number= request.POST.get('phone')
        user.country= request.POST.get('country')
        user.city= request.POST.get('city')

        user.save()

        subscriber = Subscriber()
        subscriber.subscriber_ID = user
        subscriber.password = pbkdf2_sha256.encrypt(request.POST.get('password'), rounds=12000, salt_size=23)  
        if request.FILES.get('personal_image'):
            subscriber.picture = request.FILES['personal_image']
        else:
            subscriber.picture = 'static_image/male.png' if request.POST.get('gender') == 'male' else 'static_image/female.png'
        subscriber.save()
        # here we should use a payment API to do the transition
        # if success will show a message

        subscription = Subscription()
        subscription.package_ID = Package.objects.filter(package_ID = request.POST.get('package'))[0]
        subscription.subscriber_ID = subscriber
        subscription.subscription_status_ar = subscription.subscription_status_en = 'ACTIVE'

        subscription.save()

        messages.success(request,subscribtion_payment_success[ request.session['lang'] ])
        # then send activation code to activate your account
        return send_activation_code(request, user.email)

    else:
        return redirect('mainpage')


#######
def subscriber_homepage(request):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'
        
    if(not request.session.has_key('role') or request.session['role'] != 'subscriber'):
        return redirect('mainpage')

    #calling this function to check and update the expiered ONLINE or EDITORIAL articles to ARCHIVE
    check_and_update()
    
    latest_articles = get_16_articles('latest')            # latest means latest articles
    highest_views_artilces = get_16_articles('h_views')    # h_views means highest views articles
    highest_rates_articles = get_16_articles('h_rates')    # h_rates means highest rates articles

    context = {
        'extends':'scientificJournalism/base_component/subscriber.html'
        }

    context['editorial_articles'] = Article.objects.filter(Q(is_editorial=True) & Q(article_status_en='EDITORIAL'))

    context['first_latest_articles'] = latest_articles.pop(0) if len(latest_articles) else [] 
    context['latest_articles'] = latest_articles

    context['first_highest_views_artilces'] = highest_views_artilces.pop(0) if len(highest_views_artilces) else [] 
    context['highest_views_artilces'] = highest_views_artilces


    context['first_highest_rates_articles'] = highest_rates_articles.pop(0) if len(highest_rates_articles) else [] 
    context['highest_rates_articles'] = highest_rates_articles

    context = {**context, **subscriber_base_info[ request.session['lang'] ],**home_articles_info[ request.session['lang'] ]}

    return render(request,'scientificJournalism/body_component/home_articles.html',context)
   


def error_404(request):
    return render(request,'404.html')


#######
def show_article(request,article_ID):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'
    
    #calling this function to check and update the expiered ONLINE or EDITORIAL articles to ARCHIVE
    check_and_update()
    
    try:
        article = Article.objects.get(article_ID=article_ID)
    except (Article.DoesNotExist):
        pass
    else:
        if article.article_status_en in ['ONLINE','EDITORIAL']:
            article.number_of_views += 1 if request.session.has_key('role') or not article.price else 0
            article.save()
            context ={
                'article': article,
                'writer_information' : article.writer_ID,
                'comments': Comment.objects.filter(article_ID=article.article_ID).order_by('-comment_date'),
                'average': Rate.objects.filter(article_ID=article.article_ID).aggregate(Avg('subscriber_rating')),
                'price': Setting.objects.get().article_price,
                'latest_articles':Article.objects.filter(Q(article_status_en='ONLINE') | Q(article_status_en='EDITORIAL')).order_by('-publishing_date')[0:3],
                'our_writers':Writer.objects.filter(account_status_en='ONLINE').order_by("?")[:3],
            }
            if not article.price or (request.session.has_key('role') and (request.session['role'] == 'subscriber' or (request.session['role'] == 'writer' and article.field_ar == request.session['field'] ))):
                multimedia = Multimedia.objects.filter(aritcle_ID=article.article_ID)
                context['multimedia'] = multimedia
                context['related_articles'] = Article.objects.filter(article_status_en='ONLINE',field_en=article.field_en).order_by('-number_of_views')[:3]
            
            if request.session.has_key('role') and request.session['role'] == 'subscriber':
                context = {**context, **subscriber_base_info[ request.session['lang'] ], **subscriber_show_article_info[ request.session['lang'] ] }
                return render(request,'scientificJournalism/body_component/s_article.html',context)

            if request.session.has_key('role') and request.session['role'] == 'writer':
                show_article_info[ request.session['lang'] ]['buy_article'] = show_article_info[ request.session['lang'] ]['buy_article'].format(Setting.objects.get().article_price)
                context = {**context, **writer_base_info[ request.session['lang'] ], **show_article_info[ request.session['lang'] ] }
                context['extend'] = 'scientificJournalism/base_component/writer.html'
                return render(request,'scientificJournalism/body_component/r_article.html',context)
            
            if not request.session.has_key('role'):
                show_article_info[ request.session['lang'] ]['buy_article'] = show_article_info[ request.session['lang'] ]['buy_article'].format(Setting.objects.get().article_price)
                context = {**context, **reader_base_info[ request.session['lang'] ], **show_article_info[ request.session['lang'] ] }
                context['extend'] = 'scientificJournalism/base_component/reader.html'

                # return render(request,'display.html',{'context':context['comments']})
                return render(request,'scientificJournalism/body_component/r_article.html',context)
            

    messages.error(request, article_not_exist_alert[request.session['lang']])
    return redirect('mainpage')


#######
def show_archived_article(request,article_ID):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'
    
    #calling this function to check and update the expiered ONLINE or EDITORIAL articles to ARCHIVE
    check_and_update()
    
    try:
        article = Article.objects.get(article_ID=article_ID)
    except (Article.DoesNotExist):
        pass
    else:
        if article.article_status_en == 'ARCHIVED':
            article.number_of_views += 1 #if request.session.has_key('role') or not article.price else 0
            article.save()
            context ={
                'article': article,
                'writer_information' : article.writer_ID,
                'comments': Comment.objects.filter(article_ID=article.article_ID).order_by('-comment_date'),
                'average': Rate.objects.filter(article_ID=article.article_ID).aggregate(Avg('subscriber_rating')),
                'multimedia' : Multimedia.objects.filter(aritcle_ID=article.article_ID)
                # 'price': Setting.objects.get().article_price,
                # 'latest_articles':Article.objects.filter(Q(article_status_en='ONLINE') | Q(article_status_en='EDITORIAL')).order_by('-publishing_date')[0:3],
                # 'our_writers':Writer.objects.filter(account_status_en='ONLINE').order_by("?")[:3],
            }
            # if not article.price or (request.session.has_key('role') and (request.session['role'] == 'subscriber' or (request.session['role'] == 'writer' and article.field_ar == request.session['field'] ))):
            #     multimedia = Multimedia.objects.filter(aritcle_ID=article.article_ID)
            #     context['multimedia'] = multimedia
            #     context['related_articles'] = Article.objects.filter(article_status_en='ONLINE',field_en=article.field_en).order_by('-number_of_views')[:3]
            
            if request.session.has_key('role') and request.session['role'] == 'subscriber':
                context = {**context, **subscriber_base_info[ request.session['lang'] ], **show_article_info[ request.session['lang'] ] }
                context['extend'] = 'scientificJournalism/base_component/subscriber.html'
                

            if request.session.has_key('role') and request.session['role'] == 'writer':
                context = {**context, **writer_base_info[ request.session['lang'] ], **show_article_info[ request.session['lang'] ] }
                context['extend'] = 'scientificJournalism/base_component/writer.html'

            
            if not request.session.has_key('role'):
                context = {**context, **reader_base_info[ request.session['lang'] ], **show_article_info[ request.session['lang'] ] }
                context['extend'] = 'scientificJournalism/base_component/reader.html'
            
            return render(request,'scientificJournalism/body_component/archive_article.html',context)
            
    messages.error(request, article_not_exist_alert[request.session['lang']])
    return redirect('mainpage')


#######     
def rating_article(request):
    article_id = request.GET.get('article_ID', None)
    the_rate = request.GET.get('the_rate', None)

    obj, created = Rate.objects.get_or_create(
        article_ID = Article.objects.get(article_ID = article_id),
        subscriber_ID = Subscriber.objects.get(subscriber_ID = request.session['id']),
        defaults={'subscriber_rating' : int(the_rate)}
        )
    if not created:
        obj.subscriber_rating = int(the_rate)
        obj.save()

    updated_rate = Rate.objects.filter(article_ID=article_id).aggregate(Avg('subscriber_rating')),
    data = {
        'updated_rate': round(updated_rate[0]['subscriber_rating__avg'],1),
    }
    return JsonResponse(data)


#######
def commenting_article(request):
    article_id = request.GET.get('article_ID')
    the_comment = request.GET.get('the_comment')

    Comment.objects.create(
        article_ID = Article.objects.get(article_ID = article_id),
        subscriber_ID = Subscriber.objects.get(subscriber_ID = request.session['id']),
        comment_content = the_comment
        )
    
    all_comments = Comment.objects.filter(article_ID = article_id).order_by('-comment_date')
    
    comments = []
    
    for comment in all_comments:
        single_comment = {
            'name': comment.subscriber_ID.subscriber_ID.first_name +" "+ comment.subscriber_ID.subscriber_ID.last_name,
            'image': comment.subscriber_ID.picture.url,
            'date': comment.comment_date,
            'content': comment.comment_content,
        }
        comments.append(single_comment)

    data = {
        'comments' : comments
    }

    return JsonResponse(data)


###
def archive(request):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'

    # get an appropriate header and footer template
    extend_page = ''

    if request.session.has_key('role') and request.session['role'] == 'subscriber':
        extend_page = 'scientificJournalism/base_component/subscriber.html'
    elif request.session.has_key('role') and request.session['role'] == 'writer':
        extend_page = 'scientificJournalism/base_component/writer.html'
    else:
        extend_page = 'scientificJournalism/base_component/reader.html'

    # get header and footer info
    if not request.session.has_key('role'):
        header = reader_base_info[ request.session['lang'] ]
    elif request.session['role'] == 'subscriber' : 
        header = subscriber_base_info[ request.session['lang'] ]
    elif request.session['role'] == 'writer' : 
        header = writer_base_info[ request.session['lang'] ]
    
    #calling this function to check and update the expiered ONLINE or EDITORIAL articles to ARCHIVE
    check_and_update()
    
    #all in archive status
    latest_10_archived_articles = Article.objects.filter(article_status_ar='ARCHIVED').order_by('-publishing_date')[:10]
    most_10_archived_articles_read = Article.objects.filter(article_status_ar='ARCHIVED').order_by('-number_of_views')[:10]
    highest_10_archived_articles_rate = []
    num_archived_articles_in_fields = []
    highest_writer_wrote_article = []


    # here make array of highest 10 archived articles rate
    get_highest_10_archived_articles_rate = Rate.objects.values('article_ID').annotate(rating_avg = Avg('subscriber_rating')).filter(article_ID__article_status_en='ARCHIVED').order_by('-rating_avg')[:10]

    for rated_article in get_highest_10_archived_articles_rate:
        highest_10_archived_articles_rate.append({
            'article': Article.objects.get(article_ID=rated_article['article_ID']),
            'rate': rated_article['rating_avg']
        } )
    # finish

    # here get number of archived articles in each field 
    article_fields = ['HEALTH','IT','ECOLOGY','ORGANISMS','NATURE','CHEMICAL','MATH','ENGINEERING']

    for field in article_fields:
        num_archived_articles_in_fields.append({
            'field': field,
            'num_of_articles': Article.objects.filter(Q(article_status_ar='ARCHIVED') & Q(field_ar=field)).count()
        })
    # finish

    # here get highest writers those have written articles and those artices now in archived

    wrtiers = Article.objects.values('writer_ID').annotate(articles_sum = Count('writer_ID')).filter(article_status_en='ARCHIVED').order_by('-articles_sum')[:10]

    for wrtier in wrtiers:
        the_writer = Writer.objects.get(writer_ID=wrtier['writer_ID'])
        highest_writer_wrote_article.append({
            'name': the_writer.writer_ID.first_name+' '+the_writer.writer_ID.last_name,
            'num_articles': wrtier['articles_sum']
        })


    # finish

    
    context = {
        'extend_page': extend_page,
        'latest_10_archived_articles': latest_10_archived_articles,
        'most_10_archived_articles_read': most_10_archived_articles_read,
        'highest_10_archived_articles_rate': highest_10_archived_articles_rate,
        'num_archived_articles_in_fields': num_archived_articles_in_fields,
        'highest_writer_wrote_article': highest_writer_wrote_article,
    }


    context = {**context, **header, **archive_info[ request.session['lang'] ]}
    # return render(request,'display.html',{'context':num_archived_articles_in_fields[0]})
    return render(request,'scientificJournalism/body_component/archive.html',context)



# ==================================================================================


# import os
# import shutil
# import datetime
# import convertapi
# from docx import Document
# from docx.shared import Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from zipfile import ZipFile
# from django.http import HttpResponse
# from django.conf import settings

#######
def create_docx(the_article, title):
    # create docx file include the article [title, cover_image, writer_name, publition_date, context]
    docx_file_name = title+'.docx'
    document = Document()
    
    title_format = document.add_heading(title, 0).paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_picture(the_article.cover_image, height=Inches(3.2), width=Inches(6))
    
    writer_name = the_article.writer_ID.writer_ID.first_name+' '+the_article.writer_ID.writer_ID.last_name

    writer = document.add_paragraph()
    writer_format = writer.paragraph_format
    writer_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    writer.add_run(writer_name).bold = True

    publishing_date = document.add_paragraph()
    publishing_date_format = publishing_date.paragraph_format
    publishing_date_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    publishing_date.add_run(str(the_article.publishing_date.date())).bold = True
    
    paragraph_format = document.add_paragraph(the_article.context_ar).paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.save(docx_file_name)

    return docx_file_name

# convertapi.api_secret = 'DQQN4jeXNLcICJz8'
# convertapi.api_secret = 'SK445mykUcSumgY2'
#######
def docx_to_pdf(docx_file_name,title):
    convertapi.api_secret = 'DQQN4jeXNLcICJz8'
    result = convertapi.convert('pdf', { 'File': docx_file_name })

    pdf_file_name = title+'.pdf'

    # save to file
    result.file.save('./'+pdf_file_name)

    return pdf_file_name

#######
def assign_elements_to_zipfile(pdf_file_name, article_ID, zip_file):

    zip_file.write(pdf_file_name)
    # get assotiated multimeda
    multimedia = Multimedia.objects.filter(aritcle_ID=article_ID)
    
    # loop onto multimedia and adding one after another to zip file
    for media in multimedia:
        name = media.content.url.split('/')[-1]
        source = settings.BASE_DIR+media.content.url
        distination = settings.BASE_DIR
        shutil.copy(source, distination)
        zip_file.write(name)
        os.remove(name)

#######
def buy_article(request,article_ID):

    download_code = str(datetime.datetime.now().timestamp()).replace('.','')

    the_article = Article.objects.get(article_ID = article_ID)

    title = the_article.title_ar if request.session['lang'] == 'ar' else the_article.title_en
    
    zip_file_name = '%s.zip' % download_code
    zip_file = ZipFile( zip_file_name , "w")
    
    docx_file_name = create_docx(the_article,title)

    pdf_file_name = docx_to_pdf(docx_file_name,title)

    assign_elements_to_zipfile(pdf_file_name,article_ID,zip_file)

    # close zip file that opened to write
    zip_file.close()

    #prepare to download zip flie
    zip_file =  open(zip_file_name, 'rb')
    response = HttpResponse(zip_file, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename=%s' % zip_file_name
    
    zip_file.close()

    os.remove(zip_file_name)
    os.remove(pdf_file_name)
    os.remove(docx_file_name)

    purchase = Purchase()
    purchase.article_ID = the_article
    purchase.download_code = download_code
    purchase.save()


    title_email_message = 'تمت عملية شراء المقالة بنجاح'
    email_context = 'تمت عملية شراء المقالة بنجاح.\n رمز المقالة : %s \nفي حال فقدان الملف يرجى التواصل معنا و إعطاءنا الرمز الخاص بالمقالة' % download_code
    to = request.POST.get('buyer_email')

    send_email(to, title_email_message, email_context)
    
    return response


#######
def get_field_articles(request,field):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'

    #calling this function to check and update the expiered ONLINE or EDITORIAL articles to ARCHIVE
    check_and_update()
    
    field_articles = Article.objects.filter(Q(article_status_en__in = ['ONLINE','EDITORIAL']) & Q(field_ar = field))

    articles = []

    for single_article in field_articles:
        article = {}
        article['article'] = single_article
        article['comment'] = len(Comment.objects.filter(article_ID=single_article))
        article['rate'] = Rate.objects.filter(article_ID=single_article).aggregate(Avg('subscriber_rating'))
        
        articles.append(article)
        
    extend_page = ''

    if request.session.has_key('role') and request.session['role'] == 'subscriber':
        extend_page = 'scientificJournalism/base_component/subscriber.html'
    elif request.session.has_key('role') and request.session['role'] == 'writer':
        extend_page = 'scientificJournalism/base_component/writer.html'
    else:
        extend_page = 'scientificJournalism/base_component/reader.html'

    context = {
        'extend_page': extend_page,
        'articles': articles,
        'field': field,
        'latest_articles': Article.objects.filter(Q(article_status_en='ONLINE') | Q(article_status_en='EDITORIAL')).order_by('-publishing_date')[0:3],
        'our_writers': Writer.objects.filter(account_status_en='ONLINE').order_by("?")[:3],
        'related_articles': Article.objects.filter(article_status_en='ONLINE',field_en=field).order_by('-number_of_views')[:3]
    }

    if not request.session.has_key('role'):
        header = reader_base_info[ request.session['lang'] ]
    elif request.session['role'] == 'subscriber' : 
        header = subscriber_base_info[ request.session['lang'] ]
    elif request.session['role'] == 'writer' : 
        header = writer_base_info[ request.session['lang'] ]

    context = {**context, **header, **field_articles_info[ request.session['lang'] ]}
    return render(request,'scientificJournalism/body_component/field_articles.html',context)



#######
def get_archived_field_articles(request,field):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'

    #calling this function to check and update the expiered ONLINE or EDITORIAL articles to ARCHIVE
    check_and_update()
    
    field_articles = Article.objects.filter(Q(article_status_en='ARCHIVED') & Q(field_ar = field))

    get_field = field_articles[0].get_field_ar_display if request.session['lang'] == 'ar' else field_articles[0].get_field_en_display

    articles = []

    for single_article in field_articles:
        article = {}
        article['article'] = single_article
        article['comment'] = len(Comment.objects.filter(article_ID=single_article))
        article['rate'] = Rate.objects.filter(article_ID=single_article).aggregate(Avg('subscriber_rating'))
        
        articles.append(article)
        
    extend_page = ''

    if request.session.has_key('role') and request.session['role'] == 'subscriber':
        extend_page = 'scientificJournalism/base_component/subscriber.html'
    elif request.session.has_key('role') and request.session['role'] == 'writer':
        extend_page = 'scientificJournalism/base_component/writer.html'
    else:
        extend_page = 'scientificJournalism/base_component/reader.html'

    context = {
        'extend_page': extend_page,
        'articles': articles,
        'the_field': get_field,
        # 'latest_articles': Article.objects.filter(Q(article_status_en='ONLINE') | Q(article_status_en='EDITORIAL')).order_by('-publishing_date')[0:3],
        # 'our_writers': Writer.objects.filter(account_status_en='ONLINE').order_by("?")[:3],
        # 'related_articles': Article.objects.filter(article_status_en='ONLINE',field_en=field).order_by('-number_of_views')[:3]
    }

    if not request.session.has_key('role'):
        header = reader_base_info[ request.session['lang'] ]
    elif request.session['role'] == 'subscriber' : 
        header = subscriber_base_info[ request.session['lang'] ]
    elif request.session['role'] == 'writer' : 
        header = writer_base_info[ request.session['lang'] ]

    context = {**context, **header, **field_articles_info[ request.session['lang'] ]}
    return render(request,'scientificJournalism/body_component/archived_field_articles.html',context)


######
def online_search(request):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'

    #calling this function to check and update the expiered ONLINE or EDITORIAL articles to ARCHIVE
    check_and_update()
    
    key_word = request.POST.get('key_word')

    field_articles = ''
    
    if request.session['lang'] == 'ar':
        field_articles = Article.objects.filter(Q(article_status_en__in = ['ONLINE','EDITORIAL']) & (Q(title_ar__icontains = key_word) | Q(brief_of_article_ar__icontains = key_word) | Q(writer_ID__writer_ID__first_name__icontains = key_word) | Q(writer_ID__writer_ID__last_name__icontains = key_word)))
    else:
        field_articles = Article.objects.filter(Q(article_status_en__in = ['ONLINE','EDITORIAL']) & (Q(title_en__icontains = key_word) | Q(brief_of_article_en__icontains = key_word) | Q(writer_ID__writer_ID__first_name__icontains = key_word) | Q(writer_ID__writer_ID__last_name__icontains = key_word)))

    articles = []

    for single_article in field_articles:
        article = {}
        article['article'] = single_article
        article['comment'] = len(Comment.objects.filter(article_ID=single_article))
        article['rate'] = Rate.objects.filter(article_ID=single_article).aggregate(Avg('subscriber_rating'))
        
        articles.append(article)
        
    extend_page = ''

    if request.session.has_key('role') and request.session['role'] == 'subscriber':
        extend_page = 'scientificJournalism/base_component/subscriber.html'
    elif request.session.has_key('role') and request.session['role'] == 'writer':
        extend_page = 'scientificJournalism/base_component/writer.html'
    else:
        extend_page = 'scientificJournalism/base_component/reader.html'

    context = {
        'extend_page': extend_page,
        'articles': articles,
        'latest_articles': Article.objects.filter(Q(article_status_en='ONLINE') | Q(article_status_en='EDITORIAL')).order_by('-publishing_date')[0:3],
        'our_writers': Writer.objects.filter(account_status_en='ONLINE').order_by("?")[:3],
    }

    if not request.session.has_key('role'):
        header = reader_base_info[ request.session['lang'] ]
    elif request.session['role'] == 'subscriber' : 
        header = subscriber_base_info[ request.session['lang'] ]
    elif request.session['role'] == 'writer' : 
        header = writer_base_info[ request.session['lang'] ]
    
    context = {**context, **header, **field_articles_info[ request.session['lang'] ]}

    # this also go to the same template that uses for get field_article because
    # it has the same interface
    return render(request,'scientificJournalism/body_component/field_articles.html',context)


######
def archive_search(request):
    if request.session.has_key('lang'):
        pass
    else:
        request.session['lang'] = 'ar'

    #calling this function to check and update the expiered ONLINE or EDITORIAL articles to ARCHIVE
    check_and_update()
    
    key_word = request.POST.get('archive_key_word')

    field_articles = ''
    
    if request.session['lang'] == 'ar':
        field_articles = Article.objects.filter(Q(article_status_en='ARCHIVED') & Q(title_ar__icontains = key_word) | Q(brief_of_article_ar__icontains = key_word) | Q(writer_ID__writer_ID__first_name__icontains = key_word) | Q(writer_ID__writer_ID__last_name__icontains = key_word))
    else:
        field_articles = Article.objects.filter(Q(article_status_en='ARCHIVED') & Q(title_en__icontains = key_word) | Q(brief_of_article_en__icontains = key_word) | Q(writer_ID__writer_ID__first_name__icontains = key_word) | Q(writer_ID__writer_ID__last_name__icontains = key_word))

    articles = []

    for single_article in field_articles:
        article = {}
        article['article'] = single_article
        article['comment'] = len(Comment.objects.filter(article_ID=single_article))
        article['rate'] = Rate.objects.filter(article_ID=single_article).aggregate(Avg('subscriber_rating'))
        
        articles.append(article)
        
    extend_page = ''

    if request.session.has_key('role') and request.session['role'] == 'subscriber':
        extend_page = 'scientificJournalism/base_component/subscriber.html'
    elif request.session.has_key('role') and request.session['role'] == 'writer':
        extend_page = 'scientificJournalism/base_component/writer.html'
    else:
        extend_page = 'scientificJournalism/base_component/reader.html'

    context = {
        'extend_page': extend_page,
        'articles': articles,
        'latest_articles': Article.objects.filter(article_status_en='ARCHIVED').order_by('-publishing_date')[0:3],
        'our_writers': Writer.objects.filter(account_status_en='ONLINE').order_by("?")[:3],
    }

    if not request.session.has_key('role'):
        header = reader_base_info[ request.session['lang'] ]
    elif request.session['role'] == 'subscriber' : 
        header = subscriber_base_info[ request.session['lang'] ]
    elif request.session['role'] == 'writer' : 
        header = writer_base_info[ request.session['lang'] ]
    
    context = {**context, **header, **field_articles_info[ request.session['lang'] ]}

    # this also go to the same template that uses for get field_article because
    # it has the same interface
    return render(request,'scientificJournalism/body_component/field_articles.html',context)



# ===========================================ADMIN RELETED==================================================

#######
def get_num_articles_in_each_field():
    article_fields = ['HEALTH','IT','ECOLOGY','ORGANISMS','NATURE','CHEMICAL','MATH','ENGINEERING']

    return [ Article.objects.filter(field_ar=field).count() for field in article_fields ]


#######
def get_num_articles_in_status():
    statuses = ['REJECTED','ONLINE','EDITORIAL','ARCHIVED','SUBMITTED','ASSIGNED','REVIEWED','FINAL_REVIEWED','WITH_REVIEWER','APOLOGIZED']

    return [ Article.objects.filter(article_status_ar=status).count() for status in statuses ]


#######
def view_statistics(request):
    
    num_articles_in_each_field = get_num_articles_in_each_field()
    num_articles_in_status = get_num_articles_in_status()
    context = {
        'title':'Statistics',
        'has_permission':True,
        'num_articles_in_each_field': num_articles_in_each_field,
        'num_articles_in_status': num_articles_in_status,

        }
    return render(request,'admin/statistics.html',context)


def send_email(to, title, body):
    email = EmailMessage(title, body, to=[to])
    email.send()


# ===========================================EIC==================================================

######
def EIC_homepage(request):

    if(not request.session.has_key('role') or request.session['role'] != 'EIC'):
        return redirect('mainpage')

    context = {
        'page_title':'رئيس مجلس التحرير',
        'dir':'rtl',
        'submitted_articles' : Article.objects.filter(article_status_ar="SUBMITTED"),
        'AE_list' : Associate_Editor.objects.filter(AE_ID__account_status_ar="ONLINE"),
        'articles_uneder_reviewing' : Article_Revision.objects.filter(article_ID__article_status_ar__in=["ASSIGNED","WITH_REVIEWER","REVIEWED"]),
        'reviewed_articles' : Article_Revision.objects.filter(article_ID__article_status_ar="FINAL_REVIEWED"),
    }
    # return render(request, 'display.html', {'context':context['reviewed_articles']})
    return render(request, 'scientificJournalism/body_component/EIC_homepage.html', context)


# from datetime import datetime
#######
def just_accept_article(request,article_ID):
    if(not request.session.has_key('role') or request.session['role'] != 'EIC'):
        return redirect('mainpage')

    article = Article.objects.get(article_ID=article_ID)
    feedback = Feedback()
    feedback.article_ID = article
    feedback.EIC_ID = Editor_in_Chief.objects.get(EIC_ID__EBM_ID = request.session['id'])
    feedback.writer_ID = article.writer_ID
    feedback.final_decision = 'ACCEPT'
    feedback.save()

    article.article_status_en = article.article_status_ar = "ONLINE"
    article.publishing_date = datetime.datetime.now()

    article.save()

    messages.success(request, "تم نشر المقالة بنجاح" )

    send_email(article.writer_ID.writer_ID.email,'تم قبول مقالتك','تم قبول المقالة الخاصة فيك التي بعنوان: {0} '.format(article.title_ar))

    return redirect(mainpage)


######
def accept_article_in_details(request):
    if(not request.session.has_key('role') or request.session['role'] != 'EIC'):
        return redirect('mainpage')

    priced = request.POST.get('priced')
    editorial = request.POST.get('editorial')
    article_ID = request.POST.get('article_id')
    
    article = Article.objects.get(article_ID=article_ID)

    feedback = Feedback()
    feedback.article_ID = article
    feedback.EIC_ID = Editor_in_Chief.objects.get(EIC_ID__EBM_ID = request.session['id'])
    feedback.writer_ID = article.writer_ID
    feedback.final_decision = 'ACCEPT'
    feedback.save()

    article.article_status_en = article.article_status_ar = "EDITORIAL" if editorial == 'true' else "ONLINE"
    article.publishing_date = datetime.datetime.now()
    article.price = False if priced == 'free' else True
    article.is_editorial =  True if editorial == 'true' else False

    article.save()

    messages.success(request, "تم نشر المقالة بنجاح" )

    send_email(article.writer_ID.writer_ID.email,'تم قبول مقالتك','تم قبول المقالة الخاصة فيك التي بعنوان: {0} '.format(article.title_ar))

    return redirect('mainpage')


######
def reject_article(request):
    if(not request.session.has_key('role') or request.session['role'] != 'EIC'):
        return redirect('mainpage')

    reason = request.POST.get('reason')
    article_ID = request.POST.get('article_ID')

    article = Article.objects.get(article_ID=article_ID)
    feedback = Feedback()
    feedback.article_ID = article
    feedback.EIC_ID = Editor_in_Chief.objects.get(EIC_ID__EBM_ID = request.session['id'])
    feedback.writer_ID = article.writer_ID
    feedback.EIC_comment = reason
    feedback.final_decision = 'REJECT'
    feedback.save()

    article.article_status_en = article.article_status_ar = "REJECTED"

    article.save()

    messages.success(request, "تم رفض المقالة بنجاح" )

    send_email(article.writer_ID.writer_ID.email,'تم رفض مقالتك ','تم رفض المقالة الخاصة فيك التي بعنوان: {0} \n\n سبب الرفض :- \n\n {1}'.format(article.title_ar,reason))

    return redirect(mainpage)


######
def assign_article(request):
    if(not request.session.has_key('role') or request.session['role'] != 'EIC'):
        return redirect('mainpage')

    article_id = request.POST.get('assind_article_ID')
    AE_ID= request.POST.get('assign_article')

    AE = Associate_Editor.objects.get(AE_ID__EBM_ID=AE_ID)
    article = Article.objects.get(article_ID=article_id)
    article_revision = Article_Revision()
    article_revision.EIC_ID = Editor_in_Chief.objects.get(EIC_ID__EBM_ID=request.session['id']) 
    article_revision.AE_ID = AE
    article_revision.article_ID = article

    article_revision.save()

    article.article_status_en = article.article_status_ar = "ASSIGNED"
    article.save()
    
    messages.success(request, "تم توجيه المقالة بنجاح" )


    return redirect('mainpage')



def grap_article(request,article_ID):
    if(not request.session.has_key('role') or request.session['role'] != 'EIC'):
        return redirect('mainpage')

    article = Article.objects.get(article_ID=article_ID)
    
    Article_Revision.objects.filter(article_ID=article).delete()
    
    messages.success(request, "تم سحب المقالة بنجاح" )
    
    article.article_status_en = article.article_status_ar = "SUBMITTED"
   
    article.save()

    return redirect(mainpage)



def show_EIC_article(request,article_ID):
    if(not request.session.has_key('role') or request.session['role'] != 'EIC'):
        return redirect('mainpage')

    context={
        'page_title':'رئيس مجلس التحرير',
        'dir':'rtl',
        'article': Article_Revision.objects.get(article_ID=article_ID),
        'multimedia' : Multimedia.objects.filter(aritcle_ID=article_ID),
    }
    return render(request,'scientificJournalism/body_component/EIC_article_revision.html',context)



def show_new_EIC_article(request,article_ID):
    if(not request.session.has_key('role') or request.session['role'] != 'EIC'):
        return redirect('mainpage')

    context={
        'page_title':'رئيس مجلس التحرير',
        'dir':'rtl',
        'article': Article.objects.get(article_ID=article_ID),
        'multimedia' : Multimedia.objects.filter(aritcle_ID=article_ID),
        'AE_list' : Associate_Editor.objects.filter(AE_ID__account_status_ar="ONLINE"),
    }
    # return render(request,'display.html',{'context':context['AE_list']})
    return render(request,'scientificJournalism/body_component/show_new_EIC_article.html',context)



# ===========================================AE==================================================


#######
def AE_homepage(request):
    if(not request.session.has_key('role') or request.session['role'] != 'AE'):
        return redirect('mainpage')

    context = {
        'page_title':'عضو مجلس التحرير',
        'dir':'rtl',
        'assigned_articles' : Article_Revision.objects.filter(Q(article_ID__article_status_ar="ASSIGNED") & Q(AE_ID__AE_ID__EBM_ID=request.session['id'])),
        'reviewer_list' : Reviewer.objects.filter(account_status_ar="ONLINE"),
        'articles_with_reviewer' : Article_Revision.objects.filter(article_ID__article_status_ar="WITH_REVIEWER"),
        'reviewed_articles' : Article_Revision.objects.filter(article_ID__article_status_ar="REVIEWED"),
    }

    # return render(request, 'display.html', {'context':context['reviewer_list']})
    return render(request, 'scientificJournalism/body_component/AE_homepage.html', context)


#######
def AE_apologize(request,article_ID):
    if(not request.session.has_key('role') or request.session['role'] != 'AE'):
        return redirect('mainpage')

    Article_Revision.objects.get(article_ID__article_ID=article_ID).delete()

    article = Article.objects.get(article_ID=article_ID)

    article.article_status_en = article.article_status_ar = "SUBMITTED"
    
    article.save()

    messages.success(request, "تم الإعتذار عن المقالة بنجاح" )

    return redirect('mainpage')



def show_new_AE_article(request,article_ID):
    if(not request.session.has_key('role') or request.session['role'] != 'AE'):
        return redirect('mainpage')

    context={
        'page_title':'عضو مجلس التحرير',
        'dir':'rtl',
        'article': Article.objects.get(article_ID=article_ID),
        'multimedia' : Multimedia.objects.filter(aritcle_ID=article_ID),
        'reviewer_list' : Reviewer.objects.filter(account_status_ar="ONLINE"),
    }

    return render(request,'scientificJournalism/body_component/show_new_AE_article.html',context)


#######
def invite_reviewer(request):
    if(not request.session.has_key('role') or request.session['role'] != 'AE'):
        return redirect('mainpage')

    article_id = request.POST.get('assind_article_ID')
    reviewer_ID= request.POST.get('assign_article')

    reviewer = Reviewer.objects.get(reviewer_ID=reviewer_ID)

    article = Article.objects.get(article_ID=article_id)

    article_revision = Article_Revision.objects.get(article_ID=article_id)

    article_revision.reviewer_ID = reviewer
    article_revision.invite_date = datetime.datetime.now()

    article_revision.save()

    article.article_status_en = article.article_status_ar = "WITH_REVIEWER"
    article.save()
    
    messages.success(request, "تم توجيه المقالة بنجاح" )

    send_email(reviewer.reviewer_ID.email,"دعوة لمراجعة مقالة جديدة","تمت دعوتك لمراجعة مقالة جديدة .. نرجو منك مراجعتها في أقرب وقت")

    return redirect('mainpage')



def comment_article_as_AE(request):
    if(not request.session.has_key('role') or request.session['role'] != 'AE'):
        return redirect('mainpage')

    comment = request.POST.get('comment')
    recommendation = request.POST.get('recommendation')
    article_ID = request.POST.get('article_id')
    
    article = Article.objects.get(article_ID=article_ID)

    article_revision = Article_Revision.objects.get(article_ID=article_ID)

    article_revision.AE_comment = comment
    article_revision.AE_recommendation = recommendation

    article_revision.save()


    article.article_status_en = article.article_status_ar = "FINAL_REVIEWED"

    article.save()

    messages.success(request, "تم إضافة تعليقك و توصيتك على المقالة" )

    return redirect('mainpage')



def grap_reviewer_article(request,article_ID):
    if(not request.session.has_key('role') or request.session['role'] != 'AE'):
        return redirect('mainpage')

    article = Article.objects.get(article_ID=article_ID)

    article_revision = Article_Revision.objects.get(article_ID=article)
    article_revision.invite_date = None
    article_revision.reviewer_ID = None
    
    article_revision.save()

    article.article_status_en = article.article_status_ar = "ASSIGNED"
    
    article.save()

    messages.success(request, "تم سحب المقالة من المراجع بنجاح" )

    return redirect('mainpage')



def show_AE_article(request,article_ID):
    if(not request.session.has_key('role') or request.session['role'] != 'AE'):
        return redirect('mainpage')

    context={
        'page_title':'عضو مجلس التحرير',
        'dir':'rtl',
        'article': Article_Revision.objects.get(article_ID=article_ID),
        'multimedia' : Multimedia.objects.filter(aritcle_ID=article_ID),
    }
    return render(request,'scientificJournalism/body_component/AE_article_revision.html',context)



# ===========================================Reviewer==================================================
######
def reviewer_homepage(request):
    if(not request.session.has_key('role') or request.session['role'] != 'reviewer'):
        return redirect('mainpage')

    context = {
        'page_title':'المراجع',
        'dir':'rtl',
        'reviewer_articles' : Article_Revision.objects.filter(Q(article_ID__article_status_ar="WITH_REVIEWER") & Q(reviewer_ID__reviewer_ID=request.session['id'])),
    }

    # return render(request, 'display.html', {'context':context['reviewer_list']})
    return render(request, 'scientificJournalism/body_component/reviewer_homepage.html', context)


######
def reviewer_apologize(request,article_ID):
    print(request.session['role'])
    if(not request.session.has_key('role') or request.session['role'] != 'reviewer'):
        return redirect('mainpage')
    
    article = Article.objects.get(article_ID=article_ID)

    article_revision = Article_Revision.objects.get(article_ID=article)
    article_revision.invite_date = None
    article_revision.reviewer_ID = None
    
    article_revision.save()

    article.article_status_en = article.article_status_ar = "ASSIGNED"
    
    article.save()

    messages.success(request, "تم سحب المقالة من المراجع بنجاح" )

    return redirect('mainpage')



def show_new_reviewer_article(request,article_ID):
    if(not request.session.has_key('role') or request.session['role'] != 'reviewer'):
        return redirect('mainpage')

    context={
        'page_title':'المراجع',
        'dir':'rtl',
        'article': Article.objects.get(article_ID=article_ID),
        'multimedia' : Multimedia.objects.filter(aritcle_ID=article_ID),
    }

    return render(request,'scientificJournalism/body_component/show_new_reviewer_article.html',context)



def comment_article_as_reviewer(request):
    if(not request.session.has_key('role') or request.session['role'] != 'reviewer'):
        return redirect('mainpage')

    comment = request.POST.get('comment')
    article_ID = request.POST.get('article_id')
    
    article = Article.objects.get(article_ID=article_ID)

    article_revision = Article_Revision.objects.get(article_ID=article_ID)

    article_revision.R_comment = comment

    article_revision.save()


    article.article_status_en = article.article_status_ar = "REVIEWED"

    article.save()

    messages.success(request, "تم إضافة تعليقك على المقالة" )

    return redirect('mainpage')






# ===============================================functions to update article status=============================================
# ===============================================from online or editorial to archive============================================

def check_and_update():
    """
        this function used to update articles status from
        either online or editorial to archive.
    """

    settings = Setting.objects.get()

    update = {'article_status_en': 'ARCHIVED', 'article_status_ar': 'ARCHIVED'}

    Article.objects.filter(Q(article_status_ar="ONLINE") & Q(publishing_date__lt=datetime.datetime.now()-datetime.timedelta(days=settings.online_duration))).update(**update)
    Article.objects.filter(Q(article_status_ar="EDITORIAL") & Q(publishing_date__lt=datetime.datetime.now()-datetime.timedelta(days=settings.editorial_duration))).update(**update)


# ==============================================================================================================================
# ==============================================================================================================================

